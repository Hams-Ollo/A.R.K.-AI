"""
Document processing utilities for Research Assistant.

This module handles document chunking, text cleaning, and metadata extraction
for various document formats including PDF, DOCX, TXT, MD, CSV, and Excel.
"""

import logging
from typing import List, Dict, Optional, Tuple, Any
from pathlib import Path
import re
from datetime import datetime
import pandas as pd
import fitz  # PyMuPDF
from docx import Document as DocxDocument
import markdown
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
import numpy as np
from app.utils import performance_monitor, error_tracker

# Setup logging
logger = logging.getLogger('research_assistant.processing')

class DocumentProcessor:
    """Handles multi-format document processing, chunking, and metadata extraction."""
    
    SUPPORTED_FORMATS = {
        '.pdf': 'process_pdf',
        '.docx': 'process_docx',
        '.txt': 'process_text',
        '.md': 'process_markdown',
        '.csv': 'process_tabular',
        '.xlsx': 'process_tabular',
        '.xls': 'process_tabular'
    }
    
    @error_tracker.handle_exception()
    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        min_chunk_size: int = 100
    ):
        """
        Initialize the document processor.

        Args:
            chunk_size: Target size of text chunks
            chunk_overlap: Number of characters to overlap between chunks
            min_chunk_size: Minimum size for a valid chunk
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.min_chunk_size = min_chunk_size
        
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            is_separator_regex=False
        )
        
        logger.info(f"Initialized DocumentProcessor with chunk_size={chunk_size}")
    
    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def process_document(self, file_path: str) -> Tuple[List[Document], Dict]:
        """
        Process any supported document format.

        Args:
            file_path: Path to the document file

        Returns:
            Tuple of (list of document chunks, metadata dictionary)
        """
        file_path = Path(file_path)
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
            
        file_extension = file_path.suffix.lower()
        if file_extension not in self.SUPPORTED_FORMATS:
            raise ValueError(f"Unsupported file format: {file_extension}")
            
        processor_method = getattr(self, self.SUPPORTED_FORMATS[file_extension])
        return processor_method(str(file_path))

    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def process_pdf(self, file_path: str) -> Tuple[List[Document], Dict]:
        """
        Process a PDF file and extract text and metadata.

        Args:
            file_path: Path to the PDF file

        Returns:
            Tuple of (list of document chunks, metadata dictionary)
        """
        try:
            logger.info(f"Processing PDF: {file_path}")
            pdf_path = Path(file_path)
            
            if not pdf_path.exists():
                raise FileNotFoundError(f"PDF file not found: {file_path}")
            
            # Open PDF
            pdf_document = fitz.open(file_path)
            
            # Extract metadata
            metadata = self._extract_metadata(pdf_document)
            
            # Extract and process text
            full_text = ""
            page_texts = []
            
            for page_num, page in enumerate(pdf_document):
                text = page.get_text()
                page_texts.append(text)
                full_text += f"\n{text}"
            
            # Create chunks
            chunks = self._create_chunks(full_text, metadata)
            
            # Add page numbers to chunks
            chunks = self._add_page_numbers(chunks, page_texts)
            
            logger.info(f"Successfully processed PDF with {len(chunks)} chunks")
            return chunks, metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {str(e)}")
            raise
        finally:
            if 'pdf_document' in locals():
                pdf_document.close()
    
    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def _extract_metadata(self, pdf_document) -> Dict:
        """
        Extract metadata from PDF document.

        Args:
            pdf_document: PyMuPDF document object

        Returns:
            Dictionary of metadata
        """
        try:
            metadata = pdf_document.metadata
            
            # Clean and process metadata
            processed_metadata = {
                "title": metadata.get("title", ""),
                "author": metadata.get("author", ""),
                "subject": metadata.get("subject", ""),
                "keywords": metadata.get("keywords", ""),
                "creator": metadata.get("creator", ""),
                "producer": metadata.get("producer", ""),
                "page_count": pdf_document.page_count,
                "file_size": pdf_document.stream_length,
                "format": "PDF",
                "processed_at": datetime.utcnow().isoformat()
            }
            
            # Clean empty values
            processed_metadata = {k: v for k, v in processed_metadata.items() if v}
            
            return processed_metadata
            
        except Exception as e:
            logger.error(f"Error extracting metadata: {str(e)}")
            raise
    
    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def _create_chunks(
        self,
        text: str,
        metadata: Dict,
        preserve_citations: bool = True
    ) -> List[Document]:
        """
        Create text chunks while preserving important structures.

        Args:
            text: Full document text
            metadata: Document metadata
            preserve_citations: Whether to preserve citation boundaries

        Returns:
            List of Document objects
        """
        try:
            # Clean text
            cleaned_text = self._clean_text(text)
            
            # Split text into chunks
            text_chunks = self.text_splitter.create_documents(
                texts=[cleaned_text],
                metadatas=[metadata]
            )
            
            # Post-process chunks
            processed_chunks = []
            for i, chunk in enumerate(text_chunks):
                # Add chunk metadata
                chunk.metadata.update({
                    "chunk_index": i,
                    "chunk_size": len(chunk.page_content),
                    "total_chunks": len(text_chunks)
                })
                
                processed_chunks.append(chunk)
            
            return processed_chunks
            
        except Exception as e:
            logger.error(f"Error creating chunks: {str(e)}")
            raise
    
    @error_tracker.handle_exception()
    def _clean_text(self, text: str) -> str:
        """
        Clean and normalize text content.

        Args:
            text: Input text

        Returns:
            Cleaned text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove special characters but preserve citations
        text = re.sub(r'[^\w\s\[\]\(\)\.,;:\-\'\"]+', ' ', text)
        
        # Normalize quotes
        text = re.sub(r'["""]', '"', text)
        text = re.sub(r"[']", "'", text)
        
        return text.strip()
    
    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def _add_page_numbers(
        self,
        chunks: List[Document],
        page_texts: List[str]
    ) -> List[Document]:
        """
        Add page numbers to chunks based on content matching.

        Args:
            chunks: List of document chunks
            page_texts: List of page texts

        Returns:
            Updated document chunks with page numbers
        """
        try:
            for chunk in chunks:
                chunk_text = chunk.page_content
                
                # Find the page number where this chunk appears
                for page_num, page_text in enumerate(page_texts):
                    if chunk_text in page_text:
                        chunk.metadata["page_number"] = page_num + 1
                        break
                
                # If exact match not found, use fuzzy matching
                if "page_number" not in chunk.metadata:
                    chunk.metadata["page_number"] = self._fuzzy_find_page(
                        chunk_text, page_texts
                    )
            
            return chunks
            
        except Exception as e:
            logger.error(f"Error adding page numbers: {str(e)}")
            raise
    
    @error_tracker.handle_exception()
    def _fuzzy_find_page(self, chunk_text: str, page_texts: List[str]) -> int:
        """
        Find the most likely page number using fuzzy matching.

        Args:
            chunk_text: Text to find
            page_texts: List of page texts

        Returns:
            Most likely page number (1-based)
        """
        # Simple similarity metric based on shared words
        chunk_words = set(chunk_text.split())
        max_overlap = 0
        best_page = 1
        
        for page_num, page_text in enumerate(page_texts):
            page_words = set(page_text.split())
            overlap = len(chunk_words.intersection(page_words))
            
            if overlap > max_overlap:
                max_overlap = overlap
                best_page = page_num + 1
        
        return best_page

    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def process_docx(self, file_path: str) -> Tuple[List[Document], Dict]:
        """Process a DOCX file."""
        try:
            logger.info(f"Processing DOCX: {file_path}")
            doc = DocxDocument(file_path)
            
            # Extract metadata
            metadata = {
                'title': doc.core_properties.title or Path(file_path).stem,
                'author': doc.core_properties.author,
                'created': doc.core_properties.created,
                'modified': doc.core_properties.modified,
                'file_path': file_path,
                'file_type': 'docx'
            }
            
            # Extract text with paragraph breaks
            full_text = "\n\n".join(paragraph.text for paragraph in doc.paragraphs)
            
            # Create chunks
            chunks = self._create_chunks(full_text, metadata)
            logger.info(f"Successfully processed DOCX with {len(chunks)} chunks")
            return chunks, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {str(e)}")
            raise
    
    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def process_text(self, file_path: str) -> Tuple[List[Document], Dict]:
        """Process a text file."""
        try:
            logger.info(f"Processing text file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
            
            metadata = {
                'title': Path(file_path).stem,
                'file_path': file_path,
                'file_type': 'txt',
                'created': datetime.fromtimestamp(Path(file_path).stat().st_ctime),
                'modified': datetime.fromtimestamp(Path(file_path).stat().st_mtime)
            }
            
            chunks = self._create_chunks(text, metadata)
            logger.info(f"Successfully processed text file with {len(chunks)} chunks")
            return chunks, metadata
            
        except Exception as e:
            logger.error(f"Error processing text file {file_path}: {str(e)}")
            raise
    
    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def process_markdown(self, file_path: str) -> Tuple[List[Document], Dict]:
        """Process a markdown file."""
        try:
            logger.info(f"Processing markdown file: {file_path}")
            with open(file_path, 'r', encoding='utf-8') as f:
                md_text = f.read()
            
            # Convert markdown to HTML then to plain text
            html = markdown.markdown(md_text)
            # Simple HTML tag removal (can be enhanced with BeautifulSoup if needed)
            text = re.sub('<[^<]+?>', '', html)
            
            metadata = {
                'title': Path(file_path).stem,
                'file_path': file_path,
                'file_type': 'markdown',
                'created': datetime.fromtimestamp(Path(file_path).stat().st_ctime),
                'modified': datetime.fromtimestamp(Path(file_path).stat().st_mtime)
            }
            
            chunks = self._create_chunks(text, metadata)
            logger.info(f"Successfully processed markdown file with {len(chunks)} chunks")
            return chunks, metadata
            
        except Exception as e:
            logger.error(f"Error processing markdown file {file_path}: {str(e)}")
            raise
    
    @performance_monitor.log_execution_time
    @error_tracker.handle_exception()
    def process_tabular(self, file_path: str) -> Tuple[List[Document], Dict]:
        """Process CSV or Excel files."""
        try:
            logger.info(f"Processing tabular file: {file_path}")
            file_extension = Path(file_path).suffix.lower()
            
            # Read the file with pandas
            if file_extension == '.csv':
                df = pd.read_csv(file_path)
                file_type = 'csv'
            else:
                df = pd.read_excel(file_path)
                file_type = 'excel'
            
            # Convert DataFrame to text representation
            text_chunks = []
            
            # Process each row as a chunk
            for idx, row in df.iterrows():
                # Create a text representation of the row
                row_text = f"Row {idx + 1}:\n"
                for col, value in row.items():
                    row_text += f"{col}: {value}\n"
                text_chunks.append(row_text)
            
            metadata = {
                'title': Path(file_path).stem,
                'file_path': file_path,
                'file_type': file_type,
                'columns': list(df.columns),
                'row_count': len(df),
                'created': datetime.fromtimestamp(Path(file_path).stat().st_ctime),
                'modified': datetime.fromtimestamp(Path(file_path).stat().st_mtime)
            }
            
            # Create Document objects for each chunk
            chunks = [
                Document(
                    page_content=chunk,
                    metadata={**metadata, 'chunk_id': i}
                )
                for i, chunk in enumerate(text_chunks)
            ]
            
            logger.info(f"Successfully processed tabular file with {len(chunks)} chunks")
            return chunks, metadata
            
        except Exception as e:
            logger.error(f"Error processing tabular file {file_path}: {str(e)}")
            raise
